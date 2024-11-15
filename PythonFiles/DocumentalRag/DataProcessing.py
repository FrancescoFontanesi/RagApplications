import fitz  # PyMuPDF
import re
from langchain.text_splitter import TokenTextSplitter
from transformers import AutoTokenizer
from sentence_transformers import SentenceTransformer
import psycopg
import logging
import FlexibleChunker as fc
from docx import Document


class DocumentProcessor:
    def __init__(self, model_name="efederici/sentence-BERTino", document_path=r"D:\RagApplications\Documenti\ManualeTETRAS_modificato2.pdf"):
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)
        self.document_path = document_path
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)

    def extract_subtitles_and_text(self, start_page=0):
        document = fitz.open(self.document_path)
        subtitles_dict = {}
        pattern = re.compile(r'\d+\.\d+')
        current_subtitle = None
        current_subtitle_name = None
        current_text = ""

        for page_num in range(start_page, len(document)):
            page = document.load_page(page_num)
            text = page.get_text("text")
            lines = text.split('\n')

            for i, line in enumerate(lines):
                if pattern.match(line.strip()):
                    if current_subtitle:
                        subtitles_dict[current_subtitle] = (current_subtitle_name.replace(
                            ".", ""), re.sub(r'\d+', '', current_text).strip())
                    current_subtitle = line.strip().split()[0]
                    if len(line.strip().split()) > 1:
                        current_subtitle_name = ' '.join(
                            line.strip().split()[1:])
                    else:
                        current_subtitle_name = lines[i +
                                                      1].strip() if i + 1 < len(lines) else ""
                    current_text = ""
                elif line.strip() != current_subtitle_name:
                    current_text += " " + line.strip()

            if current_subtitle:
                subtitles_dict[current_subtitle] = (current_subtitle_name.replace(
                    ".", ""), re.sub(r'\d+', '', current_text).strip())

        subtitles_dict.pop("4.2", None)
        return subtitles_dict

    def chunk_text(self, text, chunk_size=768, overlap_size=70):
        semantic_splitter = TokenTextSplitter(
            chunk_size=chunk_size,    # Number of characters per chunk
            chunk_overlap=overlap_size   # Overlap to preserve context
        )
        chunks = semantic_splitter.split_text(text)

        return chunks

    def chunk_document_plain(self, document, chunk_size=768, overlap_size=70):

        page_texts = []
        doc = Document(document)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                page_texts.append(paragraph.text)

        simple_chunks = []
        for text in page_texts:
            chunks = self.chunk_text(text, chunk_size, overlap_size)
            simple_chunks.extend(chunks)

        chunker = fc.FlexibleChunker(target_chunk_size=chunk_size)

        hybrid_chunks = []
        for text in page_texts:
            hybrid_chunk = chunker.chunk_text(text)
            hybrid_chunks.append(hybrid_chunk, [self.chunk_text(
                large_chunk, 128, 28) for large_chunk in hybrid_chunk])

        return simple_chunks, hybrid_chunks
<<<<<<< HEAD
            
    def json_to_word(self, json_file, word_file):
        # Load JSON data
        with open(json_file, 'r') as f:
            data = json.load(f)
        
        # Create a new Word document
        doc = Document()
        
        # Iterate through the dictionary and add content to the document
        for key, value in data.items():
            doc.add_heading(key, level=1)
            for item in value:
                doc.add_paragraph(item)
        
        # Save the document
        doc.save(word_file)
=======

    def generate_chunked_dictionaries(self, data_dict, chunk_size=512, overlap_percentage=[0, 0.05, 0.1, 0.15, 0.2, 0.25, 0.3]):
>>>>>>> dd1c484dd58810450947acb2b9df3234d2525e6d

        dictList = []
        for overlap in overlap_percentage:
            overlap = int(chunk_size * overlap)
            chunked_subtitles_dict = data_dict.copy()
            for subtitle, text in data_dict.items():
                if len(text[1]) <= int(chunk_size * 3.1):
                    chunked_subtitles_dict[subtitle] = [text[1]]
                else:
                    chunked_subtitles_dict[subtitle] = self.chunk_text(
                        text[1], chunk_size=chunk_size, overlap_size=overlap)
            dictList.append(chunked_subtitles_dict)
        return dictList

    def generate_chunked_dictionaries_ii(self, data_dict, chunk_size_large=512, chunk_size_list=[64, 128, 192, 256]):
        dictList = []
        for chunk_size_small in chunk_size_list:
            overlap = int(chunk_size_small * 0.2)
            chunked_subtitles_dict = data_dict.copy()
            for subtitle, text in data_dict.items():
                document = self.chunk_text(
                    text[1], chunk_size=chunk_size_large, overlap_size=0)
                chunked_subtitles_dict[subtitle] = (document, [self.chunk_text(
                    large_chunk, chunk_size_small, overlap_size=overlap) for large_chunk in document])
            dictList.append(chunked_subtitles_dict)
        return dictList

    def generate_question_dictionary(self, data_dict, chunk_size: int = 350):

        chunker = fc.FlexibleChunker(target_chunk_size=chunk_size)

        dict_for_questions = {}

        for subtitle, text in data_dict.items():
            dict_for_questions[subtitle] = chunker.chunk_text(text[1])

        for subtitle, text in dict_for_questions.items():
            if len(self.tokenizer.encode(text[-1])) < 100 and len(text) > 1:
                text[-2] += " " + text[-1]
                text.pop(-1)
        return dict_for_questions

    def get_connection_to_sql_database(self, host="localhost", data_base_name="database", username="postgres", password="postgres", port="5432"):
        g_uri = f"postgresql://{username}:{password}@{host}:{port}/{data_base_name}"
        return psycopg.connect(g_uri)

    def create_tables(self, conn):
        """Create necessary database tables using psycopg3"""
        try:
            with conn.cursor() as cur:
                cur.execute("""
        DO $$
        BEGIN
            IF NOT EXISTS (SELECT 1 FROM pg_tables WHERE tablename = 'documents') THEN
                CREATE TABLE documents (
                    id BIGINT PRIMARY KEY GENERATED BY DEFAULT AS IDENTITY,
                    content TEXT NOT NULL,
                    index TEXT,
                    embedding VECTOR(768) NOT NULL
                );
                
            END IF;
        END $$;
        """)
            conn.commit()
        except Exception as e:
            logging.error(f"Error creating tables: {e}")
            raise

    def create_index(self, conn):
        """Create index using psycopg3"""
        try:
            with conn.cursor() as cur:
                cur.execute("""
                CREATE INDEX document_embedding_idx ON documents
                USING diskann (embedding);
                """)
            conn.commit()
        except Exception as e:
            logging.error(f"Error creating index: {e}")

    def create_tables_ii(self, conn):
        """Create tables using psycopg3"""
        try:
            with conn.cursor() as cur:
                cur.execute("""
                    CREATE EXTENSION IF NOT EXISTS vector_scale CASCADE;
                    
                    CREATE TABLE IF NOT EXISTS document (
                        chunk_id BIGINT PRIMARY KEY GENERATED BY DEFAULT AS IDENTITY,
                        content TEXT,
                        index TEXT
                    );

                    CREATE TABLE IF NOT EXISTS small_chunks (
                        small_chunk_id SERIAL PRIMARY KEY,
                        large_chunk_id INTEGER REFERENCES document(chunk_id),
                        embedding vector(768)
                    );

                    
                    """)
            conn.commit()
        except Exception as e:
            logging.error(f"Error creating tables: {e}")
            raise

    def create_index_ii(self, conn):
        """Create index using psycopg3"""
        try:
            with conn.cursor() as cur:
                cur.execute("""
                CREATE INDEX small_chunk_embedding_idx ON small_chunks
                USING diskann (embedding);
                """)
            conn.commit()
        except Exception as e:
            logging.error(f"Error creating index: {e}")

    def insert_data(self, conn, data_dict: dict):
        """Insert data using psycopg3"""
        try:
            with conn.cursor() as cur:
                for index, chunks in data_dict.items():
                    embeddings = self.model.encode(chunks)

                    for chunk, embedding in zip(chunks, embeddings):

                        # Insert into document table
                        cur.execute("""INSERT INTO documents (content, index, embedding) 
                                        VALUES (%s, %s, %s)
                                        """, (chunk, index, embedding.tolist()))
            conn.commit()
        except Exception as e:
            logging.error(f"Error inserting data: {e}")
            raise

    '''
    def insert_data(self, conn, list: list[str]):
        """Insert data using psycopg3"""
        try:
                with conn.cursor() as cur:
                    embeddings = self.model.encode(list)
                        
                    for chunk, embedding in zip(list, embeddings):
                            
                        # Insert into document table
                        cur.execute("""INSERT INTO documents (content, embedding) 
                                    VALUES (%s, %s)
                                    """, (chunk, embedding.tolist()))
                conn.commit()
        except Exception as e:
            logging.error(f"Error inserting data: {e}")
            raise
'''

    def insert_data_ii(self, conn, chunked_dict: dict):
        """Insert data using psycopg3"""
        try:
            with conn.cursor() as cur:
                for index, (large_chunks, small_chunks_list) in chunked_dict.items():
                    # Process each large chunk and its small chunks
                    for large_chunk, small_chunks in zip(large_chunks, small_chunks_list):
                        # Insert main document chunk
                        cur.execute("""
                        INSERT INTO document (content, index)
                        VALUES (%s, %s)
                        RETURNING chunk_id
                        """, (large_chunk, index))

                        parent_chunk_id = cur.fetchone()[0]

                        # Create embeddings for small chunks
                        small_embeddings = self.model.encode(small_chunks)

                        # Insert small chunks with their embeddings
                        for small_chunk, embedding in zip(small_chunks, small_embeddings):
                            cur.execute("""
                            INSERT INTO small_chunks (large_chunk_id, embedding)
                            VALUES (%s, %s)
                            """, (parent_chunk_id, embedding.tolist()))

            conn.commit()
        except Exception as e:
            logging.error(f"Error inserting data: {e}")
            raise

    '''    
    def insert_data_ii(self, conn, list: list):
        """Insert data using psycopg3"""
        try:
            with conn.cursor() as cur:
                for (large_chunks, small_chunks_list) in list:
                    # Process each large chunk and its small chunks
                    for large_chunk, small_chunks in zip(large_chunks, small_chunks_list):
                        # Insert main document chunk
                        cur.execute("""
                        INSERT INTO document (content)
                        VALUES (%s, %s)
                        RETURNING chunk_id
                        """, (large_chunk))
                        
                        parent_chunk_id = cur.fetchone()[0]
                        
                        # Create embeddings for small chunks
                        small_embeddings = self.model.encode(small_chunks)
                        
                        # Insert small chunks with their embeddings
                        for small_chunk, embedding in zip(small_chunks, small_embeddings):
                            cur.execute("""
                            INSERT INTO small_chunks (large_chunk_id, embedding)
                            VALUES (%s, %s)
                            """, (parent_chunk_id, embedding.tolist()))
                            
            conn.commit()
        except Exception as e:
            logging.error(f"Error inserting data: {e}")
            raise
        

'''

    def create_view(self, conn):
        """Create a view using psycopg3"""
        try:
            with conn.cursor() as cur:
                cur.execute("""
                CREATE OR REPLACE VIEW joined_view AS
                SELECT 
                    sc.small_chunk_id,
                    sc.large_chunk_id,
                    lc.content as context,
                    lc.index as document_index,
                    sc.embedding
                FROM 
                    small_chunks sc
                JOIN document lc ON sc.large_chunk_id = lc.chunk_id;
                
                """)
                conn.commit()
        except Exception as e:
            logging.error(f"Error creating view: {e}")
            raise

    def wipe_databases(self, dbs_overlap, dbs_hybrid):
        for db in dbs_overlap:
            with db.cursor() as cur:
                cur.execute("DELETE FROM documents")
                cur.execute("DROP INDEX IF EXISTS document_embedding_idx")
                cur.execute("VACUUM")
        for db in dbs_hybrid:
            with db.cursor() as cur:
                cur.execute("DELETE FROM document")
                cur.execute("DELETE FROM small_chunks")
                cur.execute("DROP INDEX IF EXISTS small_chunk_embedding_idx")
                cur.execute("VACUUM")
<<<<<<< HEAD
    
    
    def init_for_sea(self):
        logging.info("Starting data processing...")

        subtitles_dict = self.extract_subtitles_and_text()
        logging.info("Extracted subtitles and text.")
        
        dict_list = self.generate_chunked_dictionaries(subtitles_dict, overlap_percentage=[0.3])
        logging.info("Generated chunked dictionaries.")
        
        hybrid_list = self.generate_chunked_dictionaries_ii(subtitles_dict, chunk_size_list=[128])
        logging.info("Generated hybrid chunked dictionaries.")


=======
>>>>>>> dd1c484dd58810450947acb2b9df3234d2525e6d

    def process_data(self):

        logging.basicConfig(level=logging.INFO)

        question_dict, dbs_overlap, dbs_hybrid = {}, [], []

        logging.info("Starting data processing...")

        subtitles_dict = self.extract_subtitles_and_text()
        logging.info("Extracted subtitles and text.")

        response = input(
            "Do you want to generate the chunked dictionary for the questions? (yes/no): ").strip().lower()
        if response == 'yes':
            question_dict = self.generate_question_dictionary(
                subtitles_dict, chunk_size=350)
            logging.info("Generated question dictionary.")

        response = input(
            "Do you want to generate the chunked dictionary? (yes/no): ").strip().lower()
        if response == 'yes':
            dict_list = self.generate_chunked_dictionaries(subtitles_dict)
            logging.info("Generated chunked dictionaries.")

        response = input(
            "Do you want to generate the hybrid chunked dictionary? (yes/no): ").strip().lower()
        if response == 'yes':
            hybrid_list = self.generate_chunked_dictionaries_ii(subtitles_dict)
            logging.info("Generated hybrid chunked dictionaries.")

        # The above code is creating multiple database connections to SQL databases with different
        # names and configurations. Each database connection is being established using the
        # `get_connection_to_sql_database` method with specific parameters such as host, database
        # name, username, and password. The code is setting up connections to databases named
        # "Overlap0", "Overlap0.5", "Overlap1", "Overlap1.5", "Overlap2", "Overlap2.5", "Overlap3",
        # "Hybrid64", "Hybrid128", "Hybrid192", and "Hybrid256".
        db0 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap0", username="postgres", password="578366")
        db5 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap0.5", username="postgres", password="578366")
        db10 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap1", username="postgres", password="578366")
        db15 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap1.5", username="postgres", password="578366")
        db20 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap2", username="postgres", password="578366")
        db25 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap2.5", username="postgres", password="578366")
        db30 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Overlap3", username="postgres", password="578366")
        db_hybrid_54 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Hybrid64", username="postgres", password="578366")
        db_hybrid_128 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Hybrid128", username="postgres", password="578366")
        db_hybrid_192 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Hybrid192", username="postgres", password="578366")
        db_hybrid_256 = self.get_connection_to_sql_database(
            host="localhost", data_base_name="Hybrid256", username="postgres", password="578366")

        logging.info("Connected to SQL databases.")

        dbs_overlap = [db0, db5, db10, db15, db20, db25, db30]

        response = input(
            "Do you want to add the data to the overlap dbs? (yes/no): ").strip().lower()
        if response == 'yes':
            for db, data_dict in zip(dbs_overlap, dict_list):
                self.insert_data(db, data_dict)
                self.create_index(db)
                logging.info(
                    f"Inserted data into database: Overlap {10*(dbs_overlap.index(db) * 0.5)}% ")

        dbs_hybrid = [db_hybrid_54, db_hybrid_128,
                      db_hybrid_192,  db_hybrid_256]
        response = input(
            "Do you want to add the data to the hybrid dbs? (yes/no): ").strip().lower()
        if response == 'yes':
            for db, data_dict in zip(dbs_hybrid, hybrid_list):
                self.insert_data_ii(db, data_dict)
                self.create_view(db)
                self.create_index_ii(db)
                logging.info(
                    f"Inserted hybrid data into database: {(dbs_hybrid.index(db)+1) * 64}")

        logging.info("Data processing completed.")

        return question_dict, dbs_overlap, dbs_hybrid
