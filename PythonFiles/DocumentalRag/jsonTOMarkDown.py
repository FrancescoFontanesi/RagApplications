import json
from docx import Document

def json_to_word(json_file, word_file):
    # Load JSON data
    with open(json_file, 'r') as f:
        data = json.load(f)
    
    # Create a new Word document
    doc = Document()
    
    # Iterate through the dictionary and add content to the document
    for key, value in data.items():
        doc.add_heading(key, level=1)
        for item in value:
            doc.add_paragraph(item)
    
    # Save the document
    doc.save(word_file)

if __name__ == "__main__":
    json_file = 'questions.json'  # Replace with your JSON file path
    word_file = 'questions.docx'  # Replace with your desired Word file path
    json_to_word(json_file, word_file)